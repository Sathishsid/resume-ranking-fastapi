import io
import fitz  # PyMuPDF for PDF extraction
import docx
import os
import google.generativeai as genai  # ✅ Google Gemini AI
import pandas as pd
import json
from fastapi import FastAPI, UploadFile, File, HTTPException
from typing import List
from fastapi.responses import FileResponse

# ✅ Set Your Google Gemini API Key Here
genai.configure(api_key="AIzaSyBPje4mrSr0YfllJUAcoHrqP4rCmNph-Vo")

app = FastAPI(
    title="AI-Powered Resume Ranking API",
    description="Upload job descriptions and resumes for ranking using Google Gemini.",
)

CSV_FILENAME = "resume_scores.csv"

# **Extract Text from PDF**
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extracts text from a PDF file."""
    try:
        pdf_stream = io.BytesIO(pdf_bytes)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        return "\n".join([page.get_text("text") for page in doc]).strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {str(e)}")

# **Extract Text from DOCX**
def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from DOCX: {str(e)}")

# **Extract Job Criteria Using Google Gemini**
def extract_criteria_gemini(text: str) -> List[str]:
    """Uses Google Gemini to extract key hiring criteria from a job description."""
    try:
        prompt = f"""
        Extract and list key hiring criteria from the following job description.
        Focus on required skills, experience, certifications, and qualifications.

        Job Description:
        {text}

        Return a structured JSON with key criteria in this format:
        {{
            "criteria": [
                "Must have certification XYZ",
                "5+ years of experience in Python development",
                "Strong background in Machine Learning"
            ]
        }}
        """

        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)

        # ✅ Ensure the response is valid JSON
        try:
            criteria_data = json.loads(response.text.strip())  # Convert text to dictionary
        except json.JSONDecodeError:
            print("Google Gemini returned an invalid JSON response. Trying to extract manually.")
            criteria_data = {"criteria": response.text.strip().split("\n")}

        return criteria_data.get("criteria", [])

    except Exception as e:
        print(f"Google Gemini Extraction Error: {str(e)}")
        return ["Failed to extract criteria"]

@app.post("/extract-criteria")
async def extract_criteria_api(file: UploadFile = File(...)):
    """Extracts job criteria using Google Gemini from a job description (PDF or DOCX)."""
    try:
        file_bytes = await file.read()
        file_type = file.content_type

        # Determine file type and extract text
        if file_type == "application/pdf":
            extracted_text = extract_text_from_pdf(file_bytes)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            extracted_text = extract_text_from_docx(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

        if not extracted_text:
            raise HTTPException(status_code=400, detail="No text extracted from the document.")

        # ✅ Use Google Gemini to extract criteria
        criteria = extract_criteria_gemini(extracted_text)

        return {"criteria": criteria}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract criteria: {str(e)}")
import io

# **Extract Text from PDF**
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extracts text from a PDF file."""
    try:
        pdf_stream = io.BytesIO(pdf_bytes)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        return "\n".join([page.get_text("text") for page in doc]).strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {str(e)}")

# **Extract Text from DOCX**
def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from DOCX: {str(e)}")

# **Extract Resume Details Using Google Gemini**
def extract_resume_details_gemini(text: str) -> dict:
    """Uses Google Gemini to extract skills, experience, certifications, and qualifications from a resume."""
    try:
        prompt = f"""
        Extract and list key details from the following resume.

        Resume Content:
        {text}

        Return structured JSON:
        {{
            "skills": ["list of skills"],
            "experience": "total years of experience",
            "certifications": ["list of certifications"],
            "qualifications": ["list of qualifications"]
        }}
        """

        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)

        details = response.text.strip()
        return json.loads(details)  # Ensure response is JSON

    except Exception as e:
        print(f"Google Gemini Extraction Error: {str(e)}")
        return {"skills": [], "experience": "0", "certifications": [], "qualifications": []}

# **Score Resume Using Google Gemini**
def score_with_gemini(criteria: List[str], details: dict) -> dict:
    """Uses Google Gemini to evaluate resume details against job criteria."""
    try:
        prompt = f"""
        You are an HR expert scoring resumes.

        Job Criteria:
        {', '.join(criteria)}

        Resume Details:
        Skills: {', '.join(details['skills'])}
        Experience: {details['experience']} years
        Certifications: {', '.join(details['certifications'])}
        Qualifications: {', '.join(details['qualifications'])}

        Score the resume on a scale of 0-5 for each criterion.
        Return structured JSON:
        {{
            "skills_score": "score out of 5",
            "experience_score": "score out of 5",
            "certifications_score": "score out of 5",
            "qualifications_score": "score out of 5"
        }}
        """

        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)

        scores = json.loads(response.text.strip())  # Ensure response is JSON

        # ✅ Compute final score
        total_score = sum(scores.values())

        return {
            "Skills Score": scores.get("skills_score", 0),
            "Experience Score": scores.get("experience_score", 0),
            "Certifications Score": scores.get("certifications_score", 0),
            "Qualifications Score": scores.get("qualifications_score", 0),
            "Total Score": total_score
        }

    except Exception as e:
        print(f"Google Gemini Scoring Error: {str(e)}")
        return {"Skills Score": 0, "Experience Score": 0, "Certifications Score": 0, "Qualifications Score": 0, "Total Score": 0}

@app.post("/score-resumes")
async def score_resumes(criteria: List[str], files: List[UploadFile] = File(...)):
    """Ranks resumes using Google Gemini & keyword matching."""
    if not criteria:
        raise HTTPException(status_code=400, detail="No job criteria provided.")

    scores = []

    for file in files:
        content = await file.read()
        file_type = file.content_type

        # Extract text from resume
        if file_type == "application/pdf":
            text = extract_text_from_pdf(content)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            text = extract_text_from_docx(content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.filename}")

        if not text:
            raise HTTPException(status_code=400, detail=f"No text extracted from {file.filename}")

        # ✅ Extract details using Google Gemini
        details = extract_resume_details_gemini(text)

        # ✅ Get Gemini-based score
        score_data = score_with_gemini(criteria, details)
        score_data["File Name"] = file.filename

        scores.append(score_data)

    df = pd.DataFrame(scores)

    # ✅ Save scores to CSV
    if os.path.exists(CSV_FILENAME):
        df.to_csv(CSV_FILENAME, mode="a", index=False, header=False)
    else:
        df.to_csv(CSV_FILENAME, index=False)

    return {
        "message": "Scoring completed! Data saved.",
        "csv_filename": CSV_FILENAME,
        "download_url": "/download-resume-scores",
        "scores": scores
    }

@app.get("/download-resume-scores")
async def download_resume_scores():
    """Download the generated resume scores CSV."""
    return FileResponse(CSV_FILENAME, media_type="text/csv", filename=CSV_FILENAME)
